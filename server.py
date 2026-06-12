#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""FastAPI backend for whisper-tool WebUI."""

import asyncio
import concurrent.futures
import json
import os
import re
import subprocess
import sys
import tempfile
import threading
import time
import traceback
from typing import AsyncGenerator, Optional

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

_APP_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _APP_DIR)
from app import (
    CancelledError,
    BACKEND_DEFAULT_MODELS,
    MLX_REPO_MAP,
    MODEL_LABEL_TO_NAME,
    WAV_EXT,
    _MAX_SUMMARY_INPUT_CHARS,
    detect_low_memory_default,
    extract_audio,
    get_audio_duration,
    load_history,
    load_settings,
    recommended_model_for_memory,
    save_history,
    save_settings,
    split_audio_chunks,
    summarize_text,
)

app = FastAPI(title="whisper-tool API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ========== State ==========

class _State:
    def __init__(self):
        self.staged_files: list[str] = []
        self.processing: bool = False
        self.cancel_requested: bool = False
        self.temp_files: list[str] = []
        self.batch_results: list[dict] = []
        self.batch_errors: list[dict] = []
        self.transcript_text: Optional[str] = None
        self.summary_text: Optional[str] = None
        self.output_path: Optional[str] = None
        self.summary_output_path: Optional[str] = None
        self.summary_running: bool = False
        self.summary_cancelled: bool = False
        # cached ML models
        self._diarization_pipeline = None
        self._whisper_repo: Optional[str] = None
        # job options
        self.job_model: str = "large-v3-turbo"
        self.job_diarize: bool = True
        self.job_turbo: bool = True
        self.job_lightweight: bool = False
        # SSE: 複数接続(複数タブ/端末)にブロードキャストするため購読者キューの集合を持つ
        self._subscribers: set = set()
        self._loop: Optional[asyncio.AbstractEventLoop] = None


st = _State()


def _push(event: dict):
    """Send an SSE event to ALL connected clients from any thread."""
    if not (st._loop and st._loop.is_running()):
        return
    for q in list(st._subscribers):
        asyncio.run_coroutine_threadsafe(q.put(event), st._loop)


def _resolve_output_path(source_file: str, suffix: str, settings: dict) -> str:
    """出力ファイルパスを解決する。
    構造: {output_dir}/文字起こし/{ファイル名(拡張子なし)}/{ファイル名}_{suffix}.{format}
    output_dir が空の場合は入力ファイルと同じフォルダを起点にする。
    """
    fmt = settings.get("output_format", "txt")
    base_dir = settings.get("output_dir", "").strip()
    if not base_dir:
        base_dir = os.path.dirname(source_file)

    stem = os.path.splitext(os.path.basename(source_file))[0]
    out_dir = os.path.join(os.path.expanduser(base_dir), "文字起こし", stem)
    os.makedirs(out_dir, exist_ok=True)
    return os.path.join(out_dir, f"{stem}_{suffix}.{fmt}")


def _write_output(text: str, path: str, is_summary: bool = False):
    """テキストを拡張子に応じた形式で書き出す。"""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.styles["Normal"].font.name = "Hiragino Sans"
        doc.styles["Normal"].font.size = Pt(11)
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if is_summary and stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif is_summary and stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            else:
                doc.add_paragraph(stripped)
        doc.save(path)
    else:
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)


def _check_cancel():
    if st.cancel_requested:
        raise CancelledError()


def _cleanup_temp():
    for p in list(st.temp_files):
        try:
            if os.path.exists(p):
                os.remove(p)
        except Exception:
            pass
    st.temp_files.clear()


# ========== ML pipeline ==========

def _ensure_models(model_size: str):
    repo = MLX_REPO_MAP.get(model_size, MLX_REPO_MAP["large-v3-turbo"])
    if st._whisper_repo != repo:
        _push({"type": "status", "state": "loading", "message": f"Whisperモデルを準備中: {model_size}"})
        st._whisper_repo = repo

    if st.job_diarize and st._diarization_pipeline is None:
        _push({"type": "status", "state": "loading", "message": "話者分離モデルを読込中..."})
        from pyannote.audio import Pipeline
        import torch
        hf_token_path = os.path.expanduser("~/.huggingface/token")
        hf_token = open(hf_token_path).read().strip() if os.path.exists(hf_token_path) else None
        pipeline = Pipeline.from_pretrained(
            "pyannote/speaker-diarization-community-1",
            token=hf_token,
        )
        if torch.backends.mps.is_available():
            pipeline = pipeline.to(torch.device("mps"))
        st._diarization_pipeline = pipeline


def _process_one_file(file_path: str, model_size: str) -> tuple[str, str]:
    import mlx_whisper

    fname = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    _check_cancel()

    # Step 1: WAV conversion
    if ext != WAV_EXT:
        _push({"type": "progress_detail", "message": f"{fname}: WAV変換中..."})
        fd, tmp = tempfile.mkstemp(suffix=".wav")
        os.close(fd)
        st.temp_files.append(tmp)
        extract_audio(file_path, tmp)
        audio_path = tmp
    else:
        audio_path = file_path

    _check_cancel()

    # Step 2: chunk split
    _push({"type": "progress_detail", "message": f"{fname}: 音声を分割中..."})
    chunks, _ = split_audio_chunks(audio_path, chunk_seconds=300)
    total_chunks = len(chunks)
    for cp, _ in chunks:
        if cp != audio_path:
            st.temp_files.append(cp)

    # Step 3: diarization + transcription
    diarization = None
    all_segments: list[dict] = []
    turbo = st.job_turbo and not st.job_lightweight

    def do_diarization():
        nonlocal diarization
        _push({"type": "progress_detail", "message": f"{fname}: 話者を分析中..."})
        out = st._diarization_pipeline(audio_path, num_speakers=2)
        diarization = out.speaker_diarization if hasattr(out, "speaker_diarization") else out

    def do_transcription():
        for i, (cp, offset) in enumerate(chunks):
            _check_cancel()
            pct = 18 + int((i / total_chunks) * 62)
            _push({"type": "progress_pct", "pct": pct,
                   "message": f"{fname}: 文字起こし中... {i+1}/{total_chunks}"})
            result = mlx_whisper.transcribe(
                cp,
                path_or_hf_repo=st._whisper_repo,
                language="ja",
                verbose=False,
                word_timestamps=False,
                condition_on_previous_text=False,
                no_speech_threshold=0.6,
            )
            for seg in result["segments"]:
                all_segments.append({
                    "start": seg["start"] + offset,
                    "end": seg["end"] + offset,
                    "text": seg["text"],
                })

    if turbo and st.job_diarize:
        _push({"type": "progress_detail", "message": f"{fname}: 並列処理中（高速モード）..."})
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as ex:
            d_fut = ex.submit(do_diarization)
            t_fut = ex.submit(do_transcription)
            while True:
                done, _ = concurrent.futures.wait(
                    [d_fut, t_fut], timeout=0.5,
                    return_when=concurrent.futures.ALL_COMPLETED,
                )
                if len(done) == 2:
                    break
                if st.cancel_requested:
                    _push({"type": "progress_detail", "message": "キャンセル中...話者分離の完了を待っています"})
            d_fut.result()
            t_fut.result()
    elif st.job_diarize:
        do_diarization()
        _push({"type": "progress_detail", "message": f"{fname}: 話者分離完了"})
        do_transcription()
    else:
        do_transcription()

    _check_cancel()
    _push({"type": "progress_pct", "pct": 85, "message": f"{fname}: 結合中..."})

    # Step 4: merge speaker + transcript
    if diarization is not None:
        diar_list = [
            (t.start, t.end, spk)
            for t, _, spk in diarization.itertracks(yield_label=True)
        ]
    else:
        diar_list = []

    def get_speaker(s: float, e: float) -> str:
        if not diar_list:
            return "話者1"
        best, best_ov = "不明", 0.0
        for ds, de, spk in diar_list:
            ov = max(0.0, min(e, de) - max(s, ds))
            if ov > best_ov:
                best_ov, best = ov, spk
        return best

    filler_re = re.compile(
        r'^(えー[っと]*|あー[っと]*|あのー?|まあ?|うーん|そのー?|ええと|んー+|ねえ?|うん|はい)[、。,.]?\s*'
    )

    turns: list[dict] = []
    for seg in all_segments:
        spk = get_speaker(seg["start"], seg["end"])
        text = filler_re.sub("", seg["text"].strip()).strip()
        if not text:
            continue
        if turns and turns[-1]["speaker"] == spk:
            turns[-1]["text"] += text
            turns[-1]["end"] = seg["end"]
        else:
            turns.append({"speaker": spk, "start": seg["start"], "end": seg["end"], "text": text})

    lines = []
    for t in turns:
        m, s = int(t["start"] // 60), int(t["start"] % 60)
        lines.append(f"[{m:02d}:{s:02d}] {t['speaker']}:\n{t['text']}\n")

    result_text = "\n".join(lines)

    _push({"type": "progress_pct", "pct": 95, "message": f"{fname}: 保存中..."})
    settings = load_settings()
    output_path = _resolve_output_path(file_path, "文字起こし", settings)
    _write_output(result_text, output_path, is_summary=False)

    _push({"type": "progress_pct", "pct": 100, "message": f"{fname}: 完了！"})
    return output_path, result_text


def _add_history_entry(file_path: str, output_path: str, model: str):
    history = load_history()
    entry = {
        "id": f"{int(time.time())}_{os.path.basename(file_path)}",
        "timestamp": time.strftime("%Y-%m-%d %H:%M"),
        "file": file_path,
        "output": output_path,
        "model": model,
        "summary_output": None,
    }
    history.insert(0, entry)
    save_history(history[:50])
    _push({"type": "history_updated"})


def _run_transcribe_job(model_size: str):
    st.batch_results = []
    st.batch_errors = []
    st.transcript_text = None
    start_time = time.time()
    total_files = len(st.staged_files)

    try:
        _check_cancel()
        _ensure_models(model_size)
    except CancelledError:
        _push({"type": "cancelled"})
        st.processing = False
        return
    except Exception as e:
        _push({"type": "error", "message": str(e), "detail": traceback.format_exc()})
        st.processing = False
        return

    for idx, file_path in enumerate(list(st.staged_files), start=1):
        if st.cancel_requested:
            break
        _push({"type": "progress", "current": idx, "total": total_files,
               "file": os.path.basename(file_path), "step": "transcribing"})
        try:
            output_path, result_text = _process_one_file(file_path, model_size)
            elapsed = time.time() - start_time
            st.batch_results.append({
                "file": file_path, "output": output_path,
                "elapsed": elapsed, "text": result_text,
            })
            st.transcript_text = result_text
            st.output_path = output_path
            _add_history_entry(file_path, output_path, model_size)
            _push({"type": "transcript", "text": result_text})
        except CancelledError:
            break
        except Exception as e:
            st.batch_errors.append({
                "file": file_path, "error": str(e), "detail": traceback.format_exc()
            })
            _push({"type": "file_error", "file": os.path.basename(file_path), "message": str(e)})
        finally:
            _cleanup_temp()

    batch_elapsed = time.time() - start_time
    st.processing = False

    if st.cancel_requested:
        _push({"type": "cancelled"})
    else:
        _push({
            "type": "complete",
            "elapsed": batch_elapsed,
            "success": len(st.batch_results),
            "errors": len(st.batch_errors),
            "errors_detail": st.batch_errors,
        })


def _run_summary_job(transcript: str, settings: dict):
    st.summary_running = True
    st.summary_cancelled = False
    _push({"type": "status", "state": "summarizing", "message": "サマリー生成中..."})

    def cancel_check():
        if st.cancel_requested:
            raise CancelledError()

    try:
        result = summarize_text(transcript, settings, cancel_check=cancel_check)
        st.summary_text = result
        _push({"type": "summary", "text": result})
        _push({"type": "summary_complete"})
        if st.output_path:
            # output_pathと同じフォルダ・同じ形式でサマリーを保存
            source_file = st.batch_results[-1]["file"] if st.batch_results else None
            if source_file:
                summary_path = _resolve_output_path(source_file, "サマリー", settings)
            else:
                # フォールバック: 文字起こしパスから推測
                summary_path = st.output_path.replace("_文字起こし.", "_サマリー.")
            _write_output(result, summary_path, is_summary=True)
            st.summary_output_path = summary_path
    except CancelledError:
        st.summary_cancelled = True
        _push({"type": "cancelled"})
    except Exception as e:
        _push({"type": "error", "message": str(e)})
    finally:
        st.summary_running = False


# ========== API routes ==========

# --- Files ---

class FilePathRequest(BaseModel):
    paths: list[str]


@app.get("/api/files")
def get_files():
    files = []
    for i, p in enumerate(st.staged_files):
        try:
            duration = get_audio_duration(p)
        except Exception:
            duration = 0
        files.append({"index": i, "path": p, "name": os.path.basename(p), "duration": duration})
    return {"files": files}


@app.post("/api/files")
def add_files(req: FilePathRequest):
    added = []
    for p in req.paths:
        if os.path.exists(p) and p not in st.staged_files:
            st.staged_files.append(p)
            added.append(p)
    return {"added": added, "total": len(st.staged_files)}


@app.delete("/api/files/{index}")
def remove_file(index: int):
    if index < 0 or index >= len(st.staged_files):
        raise HTTPException(404, "file not found")
    removed = st.staged_files.pop(index)
    return {"removed": removed}


@app.delete("/api/files")
def clear_files():
    st.staged_files.clear()
    return {"ok": True}


# --- Settings ---

class SettingsRequest(BaseModel):
    summary_enabled: Optional[bool] = None
    backend: Optional[str] = None
    model: Optional[str] = None
    api_key: Optional[str] = None
    output_format: Optional[str] = None  # "txt" | "md" | "docx"
    output_dir: Optional[str] = None     # "" = 入力ファイルと同じフォルダ


@app.get("/api/settings")
def get_settings():
    settings = load_settings()
    settings["model_choices"] = [
        {"label": label, "value": value}
        for label, value in [
            ("最速・精度低（tiny）", "tiny"),
            ("高速・精度低め（base）", "base"),
            ("バランス型（small）", "small"),
            ("精度高め・遅い（medium）", "medium"),
            ("高精度・速め（distil-large-v3）", "distil-large-v3"),
            ("最高精度・推奨（large-v3-turbo）", "large-v3-turbo"),
        ]
    ]
    settings["low_memory"] = detect_low_memory_default()
    settings["recommended_model"] = recommended_model_for_memory(settings["low_memory"])
    # 音声1秒あたりの処理秒数（Apple Silicon MLX実測ベース）
    settings["speed_factors"] = {
        "tiny":           0.05,
        "base":           0.07,
        "small":          0.12,
        "medium":         0.20,
        "distil-large-v3": 0.15,
        "large-v3-turbo": 0.18,
    }
    return settings


@app.post("/api/settings")
def post_settings(req: SettingsRequest):
    settings = load_settings()
    if req.summary_enabled is not None:
        settings["summary_enabled"] = req.summary_enabled
    if req.backend is not None:
        settings["backend"] = req.backend
    if req.model is not None:
        settings["model"] = req.model
    if req.api_key is not None:
        settings["api_key"] = req.api_key
    if req.output_format is not None:
        settings["output_format"] = req.output_format
    if req.output_dir is not None:
        settings["output_dir"] = req.output_dir
    save_settings(settings)
    return settings


# --- Jobs ---

class JobRequest(BaseModel):
    model: str = "large-v3-turbo"
    diarize: bool = True
    turbo: bool = True
    lightweight: bool = False


@app.post("/api/jobs")
def start_job(req: JobRequest):
    if st.processing:
        raise HTTPException(409, "処理中です")
    if not st.staged_files:
        raise HTTPException(400, "ファイルが選択されていません")

    model_size = MODEL_LABEL_TO_NAME.get(req.model, req.model)
    st.processing = True
    st.cancel_requested = False
    st.job_diarize = req.diarize
    st.job_turbo = req.turbo
    st.job_lightweight = req.lightweight

    threading.Thread(
        target=_run_transcribe_job,
        args=(model_size,),
        daemon=True,
    ).start()

    return {"ok": True, "model": model_size, "files": len(st.staged_files)}


@app.delete("/api/jobs/current")
def cancel_job():
    if not st.processing and not st.summary_running:
        raise HTTPException(400, "処理中ではありません")
    st.cancel_requested = True
    return {"ok": True}


@app.get("/api/status")
def get_status():
    return {
        "processing": st.processing,
        "summary_running": st.summary_running,
        "cancel_requested": st.cancel_requested,
        "staged_files": len(st.staged_files),
        "batch_results": len(st.batch_results),
        "batch_errors": len(st.batch_errors),
    }


# --- Results ---

@app.get("/api/transcript")
def get_transcript():
    return {"text": st.transcript_text or ""}


@app.get("/api/summary")
def get_summary():
    return {"text": st.summary_text or ""}


# --- Summary regenerate ---

class SummaryRegenRequest(BaseModel):
    transcript: Optional[str] = None


@app.post("/api/summary/regenerate")
def regenerate_summary(req: SummaryRegenRequest):
    if st.summary_running:
        raise HTTPException(409, "サマリー生成中です")
    transcript = req.transcript or st.transcript_text
    if not transcript:
        raise HTTPException(400, "文字起こしテキストがありません")
    settings = load_settings()
    st.cancel_requested = False
    threading.Thread(target=_run_summary_job, args=(transcript, settings), daemon=True).start()
    return {"ok": True}


# --- History ---

@app.get("/api/history")
def get_history_normalized():
    history = load_history()
    # 旧CTkアプリで作ったエントリはidがないので補完する
    changed = False
    for e in history:
        if not e.get("id"):
            ts = e.get("timestamp", "")
            fname = os.path.basename(e.get("file", "unknown"))
            e["id"] = f"legacy_{ts}_{fname}".replace(" ", "_").replace(":", "")
            changed = True
    if changed:
        save_history(history)
    return {"entries": history}


@app.get("/api/history/{entry_id}")
def get_history_entry(entry_id: str):
    for e in load_history():
        eid = e.get("id") or f"legacy_{e.get('timestamp','').replace(' ','_').replace(':','')}_{os.path.basename(e.get('file','unknown'))}"
        if eid == entry_id:
            result = dict(e)
            result["id"] = eid
            if e.get("output") and os.path.exists(e["output"]):
                result["transcript"] = open(e["output"], encoding="utf-8").read()
            if e.get("summary_output") and os.path.exists(e["summary_output"]):
                result["summary"] = open(e["summary_output"], encoding="utf-8").read()
            return result
    raise HTTPException(404, "entry not found")


@app.post("/api/history/{entry_id}/open")
def open_history_file(entry_id: str):
    for e in load_history():
        eid = e.get("id") or f"legacy_{e.get('timestamp','').replace(' ','_').replace(':','')}_{os.path.basename(e.get('file','unknown'))}"
        if eid == entry_id:
            path = e.get("output")
            if path and os.path.exists(path):
                subprocess.run(["open", path])
                return {"ok": True}
    raise HTTPException(404, "file not found")


# --- フォルダブラウザ ---

_AUDIO_VIDEO_EXTS = {
    ".wav", ".mp3", ".m4a", ".mp4", ".mov", ".flac", ".aac", ".mkv", ".avi", ".ogg", ".wma",
}

_DEFAULT_BROWSE_DIRS = [
    os.path.expanduser("~/Desktop"),
    os.path.expanduser("~/Downloads"),
    os.path.expanduser("~/Movies"),
    os.path.expanduser("~/Music"),
    os.path.expanduser("~/Documents"),
]


@app.get("/api/browse/scan")
def scan_files():
    """デフォルトフォルダを再帰スキャンして音声・動画ファイルを一覧返す（履歴タブ風フラットリスト）。"""
    results = []
    for base_dir in _DEFAULT_BROWSE_DIRS:
        if not os.path.isdir(base_dir):
            continue
        for root, dirs, files in os.walk(base_dir):
            # 隠しフォルダ・深すぎる階層はスキップ
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            depth = root[len(base_dir):].count(os.sep)
            if depth >= 3:
                dirs.clear()
                continue
            for fname in sorted(files):
                if fname.startswith("."):
                    continue
                ext = os.path.splitext(fname)[1].lower()
                if ext not in _AUDIO_VIDEO_EXTS:
                    continue
                full_path = os.path.join(root, fname)
                try:
                    size = os.path.getsize(full_path)
                except OSError:
                    continue
                try:
                    duration = get_audio_duration(full_path)
                except Exception:
                    duration = 0
                results.append({
                    "name": fname,
                    "path": full_path,
                    "size": size,
                    "folder": os.path.basename(root),
                    "duration": duration,
                })
    # 更新日時降順
    results.sort(key=lambda x: os.path.getmtime(x["path"]), reverse=True)
    return {"files": results}


@app.get("/api/browse")
def browse_folder(path: Optional[str] = None):
    """指定フォルダ内の音声・動画ファイルと子フォルダを返す。path省略時はデフォルト候補を返す。"""
    if path is None:
        # デフォルト: ホームディレクトリのよく使うフォルダ一覧
        return {
            "path": os.path.expanduser("~"),
            "parent": None,
            "dirs": [
                {"name": os.path.basename(d), "path": d}
                for d in _DEFAULT_BROWSE_DIRS
                if os.path.isdir(d)
            ],
            "files": [],
        }

    path = os.path.expanduser(path)
    if not os.path.isdir(path):
        raise HTTPException(400, f"フォルダが見つかりません: {path}")

    entries = os.scandir(path)
    dirs, files = [], []
    try:
        for e in sorted(entries, key=lambda x: (not x.is_dir(), x.name.lower())):
            if e.name.startswith("."):
                continue
            if e.is_dir(follow_symlinks=False):
                dirs.append({"name": e.name, "path": e.path})
            elif e.is_file():
                ext = os.path.splitext(e.name)[1].lower()
                if ext in _AUDIO_VIDEO_EXTS:
                    files.append({
                        "name": e.name,
                        "path": e.path,
                        "size": e.stat().st_size,
                    })
    except PermissionError:
        pass

    parent = str(os.path.dirname(path)) if path != os.path.expanduser("~") else None

    return {
        "path": path,
        "parent": parent,
        "dirs": dirs,
        "files": files,
    }


# --- SSE ---

@app.get("/api/events")
async def events():
    queue: asyncio.Queue = asyncio.Queue()
    st._loop = asyncio.get_event_loop()
    st._subscribers.add(queue)

    async def generate() -> AsyncGenerator[str, None]:
        try:
            while True:
                try:
                    event = await asyncio.wait_for(queue.get(), timeout=30)
                    yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
                except asyncio.TimeoutError:
                    yield 'data: {"type":"ping"}\n\n'
        except asyncio.CancelledError:
            pass
        finally:
            # 接続が切れたら購読者から外す（リーク防止）
            st._subscribers.discard(queue)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# --- Static frontend ---

_STATIC_DIR = os.path.join(_APP_DIR, "static")
if os.path.isdir(_STATIC_DIR):
    app.mount("/", StaticFiles(directory=_STATIC_DIR, html=True), name="static")


# ========== Entry point ==========

if __name__ == "__main__":
    import webbrowser
    import uvicorn

    port = 8080

    def _open_browser():
        time.sleep(1.5)
        webbrowser.open(f"http://localhost:{port}")

    threading.Thread(target=_open_browser, daemon=True).start()
    uvicorn.run(app, host="127.0.0.1", port=port, log_level="warning")
